import unittest
from datetime import date
from pathlib import Path
import tempfile
import pandas as pd
import docx

from cites_ops.reporters.status_docx_reporter import StatusDocxReporter
from cites_ops.core.stats_parser import StatsDocxParser


class TestStatusDocxReporter(unittest.TestCase):

    def setUp(self):
        self.temp_dir = tempfile.TemporaryDirectory()
        self.output_path = Path(self.temp_dir.name) / "Samadhan Setu Status_28-08-2026.docx"

        # Synthetic test dataset
        self.sample_df = pd.DataFrame([
            {"Id": "1", "Category": "Form-31", "Status": "assigned"},
            {"Id": "2", "Category": "Form-31", "Status": "resolved"},
            {"Id": "3", "Category": "Form-31", "Status": "closed"},
            {"Id": "4", "Category": "Annual_Accounts", "Status": "new"},
            {"Id": "5", "Category": "Annual_Accounts", "Status": "resolved"},
            {"Id": "6", "Category": "ECR", "Status": "feedback"},
        ])

    def tearDown(self):
        self.temp_dir.cleanup()

    def test_compute_metrics(self):
        totals, cat_df = StatusDocxReporter.compute_metrics(self.sample_df)

        self.assertEqual(totals["open"], 3)  # 1 assigned + 1 new + 1 feedback
        self.assertEqual(totals["resolved"], 2)  # 2 resolved
        self.assertEqual(totals["closed"], 1)  # 1 closed
        self.assertEqual(totals["total"], 6)

        self.assertEqual(len(cat_df), 3)  # Annual_Accounts, ECR, Form-31
        form31_row = cat_df[cat_df["Category"] == "Form-31"].iloc[0]
        self.assertEqual(form31_row["open"], 1)
        self.assertEqual(form31_row["resolved"], 1)
        self.assertEqual(form31_row["closed"], 1)
        self.assertEqual(form31_row["total"], 3)

    def test_generate_status_docx(self):
        prev_metrics = {"open": 2, "resolved": 1, "closed": 1, "total": 4}

        out_path = StatusDocxReporter.generate_status_docx(
            df=self.sample_df,
            output_path=self.output_path,
            report_date=date(2026, 8, 28),
            prev_metrics=prev_metrics,
        )

        self.assertTrue(out_path.is_file())

        # Inspect generated docx
        doc = docx.Document(str(out_path))
        self.assertTrue(any("Date: 28-08-2026" in p.text for p in doc.paragraphs))
        self.assertEqual(len(doc.tables), 2)

        # Verify Table 0
        t0 = doc.tables[0]
        self.assertEqual(len(t0.rows), 5)
        self.assertEqual(len(t0.columns), 6)
        # Row 2 (Today)
        self.assertEqual(t0.cell(2, 1).text.strip(), "28-08-2026")
        self.assertEqual(t0.cell(2, 2).text.strip(), "3")  # open
        self.assertEqual(t0.cell(2, 3).text.strip(), "2")  # resolved
        self.assertEqual(t0.cell(2, 4).text.strip(), "1")  # closed
        self.assertEqual(t0.cell(2, 5).text.strip(), "6")  # total

        # Row 4 (Difference)
        self.assertEqual(t0.cell(4, 1).text.strip(), "Difference")
        self.assertEqual(t0.cell(4, 2).text.strip(), "1")  # 3 - 2
        self.assertEqual(t0.cell(4, 3).text.strip(), "1")  # 2 - 1
        self.assertEqual(t0.cell(4, 4).text.strip(), "0")  # 1 - 1
        self.assertEqual(t0.cell(4, 5).text.strip(), "2")  # 6 - 4

        # Verify Table 1
        t1 = doc.tables[1]
        self.assertEqual(len(t1.rows), 4)  # Header + 3 categories
        self.assertEqual(t1.cell(0, 0).text.strip(), "By Category")
        self.assertEqual(t1.cell(1, 0).text.strip(), "Annual_Accounts")

        # Compact layout: no inherited paragraph gaps, smaller table type,
        # wider category column, and reduced page margins.
        category_para = t1.cell(1, 0).paragraphs[0]
        self.assertEqual(category_para.paragraph_format.space_before.pt, 0)
        self.assertEqual(category_para.paragraph_format.space_after.pt, 0)
        self.assertEqual(category_para.runs[0].font.size.pt, 9)
        self.assertGreater(t1.cell(1, 0).width.inches, t1.cell(1, 1).width.inches)
        self.assertAlmostEqual(doc.sections[0].top_margin.inches, 0.55, places=2)

        # Check that it can also be parsed by StatsDocxParser
        parsed = StatsDocxParser.parse_file(out_path)
        self.assertIsNotNone(parsed)
        self.assertEqual(parsed["totals"]["open"], 3)
        self.assertEqual(parsed["totals"]["resolved"], 2)
        self.assertEqual(parsed["totals"]["total"], 6)


if __name__ == "__main__":
    unittest.main()
