from datetime import date
from pathlib import Path
from typing import Any, Dict, List, Optional, Union
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

class DocxReporter:
    """
    Generates official Word documents (.docx) for administrative notes
    and sanitized Field Office Knowledge Notes.
    """

    @classmethod
    def generate_govt_note(
        cls,
        major_issues: List[Dict[str, str]],
        output_path: Union[str, Path],
        department_name: str = "EMPLOYEES' PROVIDENT FUND ORGANISATION\n(MINISTRY OF LABOUR & EMPLOYMENT, GOVT. OF INDIA)\nHEAD OFFICE",
        subject: str = "Major Issue Categories & Systemic Discrepancies - Reg.",
        note_date: Optional[Union[str, date]] = None,
    ) -> str:
        out_file = Path(output_path)
        out_file.parent.mkdir(parents=True, exist_ok=True)
        doc = docx.Document()

        # Styles
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Times New Roman"
        font.size = Pt(12)

        # Header Title
        title_heading = doc.add_heading("NOTE FOR INFORMATION & ACTION", 0)
        title_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Department
        p_dep = doc.add_paragraph()
        p_dep.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_dep.add_run(department_name)
        run.bold = True
        run.font.size = Pt(13)

        # Date
        d_str = str(note_date or date.today().strftime("%d-%m-%Y"))
        p_date = doc.add_paragraph()
        p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_date.add_run(f"Date: {d_str}")

        # Subject
        p_sub = doc.add_paragraph()
        p_sub.add_run("Subject: ").bold = True
        p_sub.add_run(subject).bold = True

        # Introduction
        p_intro = doc.add_paragraph()
        p_intro.add_run("1. ").bold = True
        p_intro.add_run(
            "A comprehensive review of the operational issue tracker and field office escalations "
            "has been conducted. Based on the observations recorded, critical systemic anomalies and "
            "processing bottlenecks have been identified requiring immediate intervention."
        )

        # Body
        p_body = doc.add_paragraph()
        p_body.add_run("2. ").bold = True
        p_body.add_run("The major issue categories reported are summarized below:")

        for idx, item in enumerate(major_issues, 1):
            title = item.get("title", f"Issue Category {idx}")
            desc = item.get("description", "")
            p_item = doc.add_paragraph()
            p_item.add_run(f"   ({chr(64 + idx)}) {title}\n").bold = True
            p_item.add_run(f"   {desc}")

        doc.save(out_file)
        return str(out_file)

    @classmethod
    def generate_knowledge_note(
        cls,
        knowledge_items: List[Dict[str, Any]],
        output_path: Union[str, Path],
        note_date: Optional[Union[str, date]] = None,
    ) -> str:
        out_file = Path(output_path)
        out_file.parent.mkdir(parents=True, exist_ok=True)
        doc = docx.Document()

        # Styles
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)

        d_str = str(note_date or date.today())
        title_p = doc.add_heading("Field Office Operations Knowledge Note", 0)

        p_sub = doc.add_paragraph()
        p_sub.add_run(f"Date: {d_str} · Technical Guidance Extracted from Operational Support Discussions\n").italic = True
        p_sub.add_run("Note: Sensitive personal and member identifiers have been masked automatically.").italic = True

        if not knowledge_items:
            doc.add_paragraph("No technical solution items were recorded for this period.")
        else:
            for idx, item in enumerate(knowledge_items, 1):
                p_item = doc.add_paragraph()
                p_item.add_run(f"{idx}. Guidance / Resolution from {item.get('sender', 'Technical Team')}\n").bold = True
                p_item.add_run(f"   {item.get('sanitized_message', '')}\n")
                if item.get("linked_issues"):
                    p_item.add_run(f"   [Cross-referenced Tickets: {', '.join(item['linked_issues'])}]\n").italic = True

        doc.save(out_file)
        return str(out_file)
