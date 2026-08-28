"""Reporter for generating Samadhan Setu Daily Status Word (.docx) documents."""

from __future__ import annotations

from datetime import datetime, date, timedelta
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple, Union
import pandas as pd
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

from ..core.stats_parser import StatsDocxParser


class StatusDocxReporter:
    """
    Generates official daily Samadhan Setu Status Word reports matching the standard enterprise template.
    """

    @classmethod
    def set_cell_margins(cls, cell, top=100, bottom=100, left=150, right=150):
        """Sets internal cell margins (padding) in dxa units."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for margin_name, val in (('top', top), ('bottom', bottom), ('left', left), ('right', right)):
            node = OxmlElement(f'w:{margin_name}')
            node.set(qn('w:w'), str(val))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
        tcPr.append(tcMar)

    @classmethod
    def set_cell_shading(cls, cell, color_hex="F2F2F2"):
        """Applies background shading to a table cell."""
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
        cell._tc.get_or_add_tcPr().append(shading_elm)

    @classmethod
    def set_table_borders(cls, table):
        """Applies standard thin black borders around all table cells."""
        tblPr = table._tbl.tblPr
        borders_elm = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>\n'
            f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>\n'
            f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>\n'
            f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>\n'
            f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>\n'
            f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>\n'
            f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>\n'
            f'</w:tblBorders>'
        )
        tblPr.append(borders_elm)

    @classmethod
    def format_cell_paragraph(cls, cell, alignment):
        """Apply compact, readable formatting without inherited paragraph gaps."""
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = alignment
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        return paragraph

    @classmethod
    def set_column_widths(cls, table, widths_inches):
        """Set deliberate widths so category names do not wrap unnecessarily."""
        table.autofit = False
        for col_idx, width_inches in enumerate(widths_inches):
            width = Inches(width_inches)
            table.columns[col_idx].width = width
            for row in table.rows:
                row.cells[col_idx].width = width

    @classmethod
    def compute_metrics(cls, df: pd.DataFrame) -> Tuple[Dict[str, int], pd.DataFrame]:
        """
        Calculates overall totals and category-wise open, resolved, closed breakdown.
        """
        if df.empty:
            empty_totals = {"open": 0, "resolved": 0, "closed": 0, "total": 0}
            empty_cat = pd.DataFrame(columns=["Category", "open", "resolved", "closed", "total"])
            return empty_totals, empty_cat

        df_work = df.copy()
        # Normalize status
        def categorize_status(s: Any) -> str:
            val = str(s or "").strip().lower()
            if val in ("resolved", "fixed"):
                return "resolved"
            elif val in ("closed",):
                return "closed"
            else:
                return "open"

        df_work["status_group"] = df_work["Status"].apply(categorize_status)
        df_work["Category"] = df_work["Category"].fillna("General").astype(str).str.strip()
        df_work["Category"] = df_work["Category"].replace("", "General")

        # Pivot category counts
        pivot = pd.pivot_table(
            df_work,
            index="Category",
            columns="status_group",
            values="Id" if "Id" in df_work.columns else df_work.columns[0],
            aggfunc="count",
            fill_value=0,
        )

        for col in ["open", "resolved", "closed"]:
            if col not in pivot.columns:
                pivot[col] = 0

        pivot["total"] = pivot["open"] + pivot["resolved"] + pivot["closed"]
        pivot = pivot.sort_index()
        pivot = pivot.reset_index()

        total_open = int(pivot["open"].sum())
        total_resolved = int(pivot["resolved"].sum())
        total_closed = int(pivot["closed"].sum())
        total_all = int(pivot["total"].sum())

        totals = {
            "open": total_open,
            "resolved": total_resolved,
            "closed": total_closed,
            "total": total_all,
        }

        return totals, pivot

    @classmethod
    def generate_status_docx(
        cls,
        df: pd.DataFrame,
        output_path: Union[str, Path],
        report_date: Optional[Union[date, str]] = None,
        prev_stats_docx: Optional[Union[str, Path]] = None,
        prev_metrics: Optional[Dict[str, Any]] = None,
    ) -> Path:
        """
        Builds the Samadhan Setu Status DOCX matching the exact format of reference files.
        """
        out_path = Path(output_path)
        out_path.parent.mkdir(parents=True, exist_ok=True)

        # Determine report date
        if report_date is None:
            rep_date = date.today()
        elif isinstance(report_date, str):
            try:
                rep_date = datetime.strptime(report_date, "%Y-%m-%d").date()
            except ValueError:
                try:
                    rep_date = datetime.strptime(report_date, "%d-%m-%Y").date()
                except ValueError as exc:
                    raise ValueError(
                        f"Invalid report date {report_date!r}; expected YYYY-MM-DD or DD-MM-YYYY"
                    ) from exc
        else:
            rep_date = report_date

        date_str_formatted = rep_date.strftime("%d-%m-%Y")
        prev_date_str = (rep_date - timedelta(days=1)).strftime("%d-%m-%Y")

        totals_today, cat_df = cls.compute_metrics(df)

        # Parse previous stats if provided
        prev_totals = {"open": 0, "resolved": 0, "closed": 0, "total": 0}
        if prev_metrics:
            prev_totals.update(prev_metrics)
        elif prev_stats_docx and Path(prev_stats_docx).is_file():
            parsed = StatsDocxParser.parse_file(prev_stats_docx)
            if parsed and "totals" in parsed:
                prev_totals = parsed["totals"]
                if "source" in parsed and "data_date" in parsed["source"]:
                    try:
                        p_dt = datetime.fromisoformat(parsed["source"]["data_date"]).date()
                        prev_date_str = p_dt.strftime("%d-%m-%Y")
                    except Exception:
                        pass
        else:
            # If no previous file given, default to today's totals to avoid negative difference
            prev_totals = {
                "open": totals_today["open"],
                "resolved": totals_today["resolved"],
                "closed": totals_today["closed"],
                "total": totals_today["total"],
            }

        diff_open = totals_today["open"] - prev_totals["open"]
        diff_resolved = totals_today["resolved"] - prev_totals["resolved"]
        diff_closed = totals_today["closed"] - prev_totals["closed"]
        diff_total = totals_today["total"] - prev_totals["total"]

        # Create Word Document
        doc = docx.Document()

        # Page margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.55)
            section.bottom_margin = Inches(0.55)
            section.left_margin = Inches(0.6)
            section.right_margin = Inches(0.6)

        # Paragraph 1: Date header
        p_date = doc.add_paragraph()
        p_date.paragraph_format.space_before = Pt(0)
        p_date.paragraph_format.space_after = Pt(3)
        r_date = p_date.add_run(f"Date: {date_str_formatted}")
        r_date.font.name = "Calibri"
        r_date.font.size = Pt(10)
        r_date.font.bold = False

        # Paragraph 2: Intro
        p_intro = doc.add_paragraph()
        p_intro.paragraph_format.space_before = Pt(0)
        p_intro.paragraph_format.space_after = Pt(5)
        r_intro = p_intro.add_run("Please refer the below table, wherein the progress of resolving/closed status of Samadhan Setu is.")
        r_intro.font.name = "Calibri"
        r_intro.font.size = Pt(10)

        # =========================================================================
        # Table 0: Progress Status (5 rows, 6 cols)
        # =========================================================================
        t0 = doc.add_table(rows=5, cols=6)
        t0.alignment = WD_TABLE_ALIGNMENT.CENTER
        cls.set_table_borders(t0)
        cls.set_column_widths(t0, [0.55, 1.65, 1.2, 1.2, 1.2, 1.2])

        # Row 0: Merged Header
        cell_00 = t0.cell(0, 0)
        cell_05 = t0.cell(0, 5)
        merged_0 = cell_00.merge(cell_05)
        p_hdr = cls.format_cell_paragraph(merged_0, WD_PARAGRAPH_ALIGNMENT.CENTER)
        r_hdr = p_hdr.add_run("Samadhan Setu Progress Status")
        r_hdr.font.name = "Calibri"
        r_hdr.font.size = Pt(9.5)
        r_hdr.font.bold = True
        cls.set_cell_margins(merged_0, top=40, bottom=40, left=80, right=80)

        # Row 1: Column Titles
        col_titles = ["S.no", "Date", "Open", "Resolved", "Closed", "Total"]
        for idx, title in enumerate(col_titles):
            cell = t0.cell(1, idx)
            p = cls.format_cell_paragraph(cell, WD_PARAGRAPH_ALIGNMENT.CENTER)
            r = p.add_run(title)
            r.font.name = "Calibri"
            r.font.size = Pt(9)
            r.font.bold = True
            cls.set_cell_margins(cell, top=30, bottom=30, left=60, right=60)

        # Row 2: Today (S.no 2)
        row2_data = ["2", date_str_formatted, str(totals_today["open"]), str(totals_today["resolved"]), str(totals_today["closed"]), str(totals_today["total"])]
        for idx, val in enumerate(row2_data):
            cell = t0.cell(2, idx)
            p = cls.format_cell_paragraph(cell, WD_PARAGRAPH_ALIGNMENT.CENTER)
            r = p.add_run(val)
            r.font.name = "Calibri"
            r.font.size = Pt(9)
            cls.set_cell_margins(cell, top=20, bottom=20, left=60, right=60)

        # Row 3: Yesterday / Previous (S.no 1)
        row3_data = ["1", prev_date_str, str(prev_totals["open"]), str(prev_totals["resolved"]), str(prev_totals["closed"]), str(prev_totals["total"])]
        for idx, val in enumerate(row3_data):
            cell = t0.cell(3, idx)
            p = cls.format_cell_paragraph(cell, WD_PARAGRAPH_ALIGNMENT.CENTER)
            r = p.add_run(val)
            r.font.name = "Calibri"
            r.font.size = Pt(9)
            cls.set_cell_margins(cell, top=20, bottom=20, left=60, right=60)

        # Row 4: Difference
        row4_data = ["", "Difference", str(diff_open), str(diff_resolved), str(diff_closed), str(diff_total)]
        for idx, val in enumerate(row4_data):
            cell = t0.cell(4, idx)
            p = cls.format_cell_paragraph(cell, WD_PARAGRAPH_ALIGNMENT.CENTER)
            r = p.add_run(val)
            r.font.name = "Calibri"
            r.font.size = Pt(9)
            if idx == 1:
                r.font.bold = True
            cls.set_cell_margins(cell, top=20, bottom=20, left=60, right=60)

        # Paragraph 3: Category intro
        p_cat_intro = doc.add_paragraph()
        p_cat_intro.paragraph_format.space_before = Pt(6)
        p_cat_intro.paragraph_format.space_after = Pt(4)
        r_cat_intro = p_cat_intro.add_run("Further, the category wise status is also given as under:")
        r_cat_intro.font.name = "Calibri"
        r_cat_intro.font.size = Pt(10)

        # =========================================================================
        # Table 1: Category Wise Status (N+1 rows, 5 cols)
        # =========================================================================
        num_cat_rows = len(cat_df) + 1
        t1 = doc.add_table(rows=num_cat_rows, cols=5)
        t1.alignment = WD_TABLE_ALIGNMENT.CENTER
        cls.set_table_borders(t1)
        cls.set_column_widths(t1, [3.7, 0.9, 0.9, 0.9, 0.9])

        # Table 1 Header
        t1_headers = ["By Category", "open", "resolved", "closed", "total"]
        for idx, h_text in enumerate(t1_headers):
            cell = t1.cell(0, idx)
            alignment = WD_PARAGRAPH_ALIGNMENT.LEFT if idx == 0 else WD_PARAGRAPH_ALIGNMENT.CENTER
            p = cls.format_cell_paragraph(cell, alignment)
            r = p.add_run(h_text)
            r.font.name = "Calibri"
            r.font.size = Pt(9)
            r.font.bold = True
            cls.set_cell_margins(cell, top=30, bottom=30, left=60, right=60)

        # Table 1 Category Rows
        for r_idx, row in cat_df.iterrows():
            table_row_idx = r_idx + 1
            row_vals = [
                str(row["Category"]),
                str(int(row["open"])),
                str(int(row["resolved"])),
                str(int(row["closed"])),
                str(int(row["total"])),
            ]
            for col_idx, val in enumerate(row_vals):
                cell = t1.cell(table_row_idx, col_idx)
                alignment = WD_PARAGRAPH_ALIGNMENT.LEFT if col_idx == 0 else WD_PARAGRAPH_ALIGNMENT.CENTER
                p = cls.format_cell_paragraph(cell, alignment)
                r = p.add_run(val)
                r.font.name = "Calibri"
                r.font.size = Pt(9)
                cls.set_cell_margins(cell, top=15, bottom=15, left=60, right=60)

        doc.save(str(out_path))
        return out_path
